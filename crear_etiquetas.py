import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, Mm
import os

def crear_etiquetas_por_cantidad(nombres, cantidad):
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.size = Pt(14)
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)

    etiquetas = ''
    if (len(nombres) < 20):
        frec = 3
    elif (len(nombres) > 20 and len(nombres) < 25):
        frec = 2
    else:
        frec = 1
    contador = 0
    for i in range(cantidad):
        if (contador <= frec):
            etiquetas += f"{nombres}     "
            contador += 1
        else:
            contador = 0
            etiquetas += '\n'
            etiquetas += f"{nombres}     "
            contador += 1

    paragraph = document.add_paragraph(etiquetas)
    paragraph.space_after = Mm(0)

    section = document.sections[0]
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    paragraph.paragraph_format.left_indent = Mm(5)
    paragraph.paragraph_format.right_indent = Mm(5)

    file_name = 'etiquetas_por_hojas.docx'
    document.save(file_name)
    abrir_archivo(file_name)

def abrir_archivo(nombre_archivo):
    os.system(nombre_archivo)

def crear_etiquetas_por_cantidad_de_hojas(nombres, cantidad_hojas):
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.size = Pt(14)
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)

    etiquetas = ''
    cantidad_hojas = cantidad_hojas * 41
    if (len(nombres) < 20):
        frec = 3
    elif (len(nombres) > 20 and len(nombres) < 25):
        frec = 2
    else:
        frec = 1
    for i in range(1, cantidad_hojas):
        contador = 0
        while(contador <= frec):
            etiquetas += f"{nombres}     "
            contador += 1
        etiquetas += '\n'

    paragraph = document.add_paragraph(etiquetas)
    paragraph.space_after = Mm(0)
    section = document.sections[0]
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.right_indent = Mm(0)
    file_name = 'etiquetas_por_cantidad.docx'
    document.save(file_name)
    abrir_archivo(file_name)

def crear_app():
    def on_click_cantidad():
        nombres = entry_nombres.get()
        cantidad = int(entry_cantidad.get())
        crear_etiquetas_por_cantidad(nombres, cantidad)

    def on_click_hojas():
        nombres = entry_nombres.get()
        cantidad_hojas = int(entry_cantidad_hojas.get())
        crear_etiquetas_por_cantidad_de_hojas(nombres, cantidad_hojas)

    root = tk.Tk()
    root.title("Generador de Etiquetas")

    label_nombres = tk.Label(root, text="Nombres:")
    label_nombres.grid(row=0, column=0, padx=5, pady=5)
    entry_nombres = tk.Entry(root, width=50)
    entry_nombres.grid(row=0, column=1, columnspan=2, padx=5, pady=5)

    label_cantidad = tk.Label(root, text="Cantidad por nombre:")
    label_cantidad.grid(row=1, column=0, padx=5, pady=5)
    entry_cantidad = tk.Entry(root, width=10)
    entry_cantidad.grid(row=1, column=1, padx=5, pady=5)

    button_cantidad = tk.Button(root, text="Crear por cantidad", command=on_click_cantidad)
    button_cantidad.grid(row=1, column=2, padx=5, pady=5)

    label_cantidad_hojas = tk.Label(root, text="Cantidad de hojas:")
    label_cantidad_hojas.grid(row=2, column=0, padx=5, pady=5)
    entry_cantidad_hojas = tk.Entry(root, width=10)
    entry_cantidad_hojas.grid(row=2, column=1, padx=5, pady=5)

    button_hojas = tk.Button(root, text="Crear por cantidad de hojas", command=on_click_hojas)
    button_hojas.grid(row=2, column=2, padx=5, pady=5)

    root.mainloop()

if __name__ == "__main__":
    crear_app()
